import streamlit as st
import fitz  # PyMuPDF
from docx import Document
from io import BytesIO

# --- CẤU HÌNH TRANG WEB ---
st.set_page_config(page_title="PDF Extraction Tool", layout="wide")
st.title("📄 Tool Trích Xuất PDF 2 Cột -> Word")
st.markdown("Công cụ chuyên trị các file PDF chia 2 cột (Paper, Báo cáo).")

# --- HÀM XỬ LÝ (GIỮ NGUYÊN LOGIC CŨ) ---
@st.cache_data # Cache giúp không phải load lại PDF mỗi lần tìm từ khóa mới
def process_pdf(uploaded_file):
    doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    sorted_text = []
    
    for page in doc:
        width = page.rect.width
        mid_x = width / 2
        
        blocks = page.get_text("blocks")
        col_left = []
        col_right = []
        
        for b in blocks:
            if b[6] != 0: continue # Bỏ qua ảnh
            if b[0] < mid_x:
                col_left.append(b)
            else:
                col_right.append(b)
        
        col_left.sort(key=lambda x: x[1])
        col_right.sort(key=lambda x: x[1])
        
        ordered_blocks = col_left + col_right
        
        for b in ordered_blocks:
            text = b[4].strip()
            if text:
                sorted_text.append(text)
    return sorted_text, doc.name

# --- GIAO DIỆN CHÍNH ---
col1, col2 = st.columns([1, 2])

with col1:
    st.header("1. Tải file")
    uploaded_file = st.file_uploader("Chọn file PDF", type="pdf")

if uploaded_file is not None:
    # Xử lý file ngay khi tải lên
    try:
        all_text, filename = process_pdf(uploaded_file)
        st.success(f"✅ Đã đọc xong file! Tổng {len(all_text)} đoạn văn.")
        
        with col2:
            st.header("2. Tìm kiếm & Xuất")
            
            with st.form("search_form"):
                keyword = st.text_input("Nhập từ khóa bắt đầu:")
                num_paras = st.number_input("Số đoạn văn muốn lấy:", min_value=1, value=5)
                ignore_case = st.checkbox("Không phân biệt hoa/thường", value=True)
                
                submitted = st.form_submit_button("🚀 Trích xuất ngay")
                
            if submitted and keyword:
                # Logic tìm kiếm
                start_idx = -1
                target = keyword.lower() if ignore_case else keyword
                
                for i, text in enumerate(all_text):
                    check_text = text.lower() if ignore_case else text
                    if target in check_text:
                        start_idx = i
                        break
                
                if start_idx == -1:
                    st.error(f"❌ Không tìm thấy từ khóa '{keyword}'")
                else:
                    # Logic lấy đoạn văn
                    result_paras = []
                    first_para = all_text[start_idx]
                    pos = (first_para.lower() if ignore_case else first_para).find(target)
                    result_paras.append(first_para[pos:])
                    result_paras.extend(all_text[start_idx+1 : start_idx+num_paras])
                    
                    # Hiển thị xem trước
                    st.info("Kết quả tìm thấy:")
                    preview_text = "\n\n--- NGẮT ---\n\n".join(result_paras)
                    st.text_area("Preview", preview_text, height=300)
                    
                    # Tạo file Word trong bộ nhớ (không lưu ra đĩa server)
                    doc = Document()
                    doc.add_heading(f'Trích xuất: "{keyword}"', 0)
                    doc.add_paragraph(f"Nguồn file: {uploaded_file.name}")
                    for p in result_paras:
                        doc.add_paragraph(p)
                        doc.add_paragraph("---")
                    
                    # Chuẩn bị file để tải về
                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    
                    st.download_button(
                        label="⬇️ Tải file Word (.docx)",
                        data=buffer,
                        file_name=f"KetQua_{keyword}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
    except Exception as e:
        st.error(f"Có lỗi xảy ra: {e}")

else:
    with col2:
        st.info("👈 Vui lòng tải file PDF bên cột trái trước.")